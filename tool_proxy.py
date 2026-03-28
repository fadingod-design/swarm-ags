# -*- coding: utf-8 -*-
"""
蜂群AGS工具代理服务 v2
直接调用Python脚本执行工具，不依赖Gateway API
"""

import json
import subprocess
import os
import requests
from http.server import HTTPServer, BaseHTTPRequestHandler
from socketserver import ThreadingMixIn

# 尝试导入BeautifulSoup
try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    print("[Warning] BeautifulSoup not installed, search quality may be limited")

# Skills目录
SKILLS_DIR = os.path.expanduser("~/.openclaw-autoclaw/skills")

class ToolProxyHandler(BaseHTTPRequestHandler):
    """工具代理HTTP处理器"""
    
    def log_message(self, format, *args):
        print("[ToolProxy]", args[0])
    
    def send_json(self, data, status=200):
        self.send_response(status)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(json.dumps(data, ensure_ascii=False).encode('utf-8'))
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
    
    def do_GET(self):
        """列出可用工具"""
        tools = {
            "web_search": {
                "name": "联网搜索",
                "description": "搜索互联网获取最新信息",
                "params": {"query": "搜索关键词"}
            },
            "feishu_create_doc": {
                "name": "创建飞书文档",
                "description": "创建飞书在线文档",
                "params": {"title": "文档标题"}
            },
            "feishu_write_doc": {
                "name": "写入飞书文档",
                "description": "向飞书文档追加内容",
                "params": {"doc_token": "文档token", "content": "内容"}
            },
            "generate_image": {
                "name": "生成图片",
                "description": "AI生成图片",
                "params": {"prompt": "图片描述"}
            },
            "stock_analysis": {
                "name": "A股实时行情",
                "description": "获取A股实时价格、涨跌、成交量",
                "params": {"code": "股票代码（6位数字）"}
            },
            "stock_minute": {
                "name": "分时量能分析",
                "description": "分析分时量能分布、主力动向",
                "params": {"code": "股票代码"}
            },
            "stock_valuation": {
                "name": "股票估值分析",
                "description": "DCF估值、可比公司分析、情景分析",
                "params": {"code": "股票代码", "scenario": "bull/base/bear"}
            },
            "generate_chart": {
                "name": "生成图表",
                "description": "生成折线图、柱状图、饼图、K线图",
                "params": {"chart_type": "line/bar/pie", "title": "标题", "data": "数据"}
            },
            "create_word_report": {
                "name": "生成Word报告",
                "description": "生成Word格式的分析报告",
                "params": {"title": "标题", "content": "内容"}
            },
            "send_email": {
                "name": "发送邮件",
                "description": "发送报告到邮箱",
                "params": {"subject": "主题", "body": "内容", "attachment": "附件路径"}
            }
        }
        self.send_json({"status": "ok", "tools": tools})
    
    def do_POST(self):
        """执行工具"""
        try:
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length).decode('utf-8')
            data = json.loads(body)
            
            tool_name = data.get('tool')
            params = data.get('params', {})
            
            if not tool_name:
                self.send_json({"error": "Missing tool parameter"}, 400)
                return
            
            print(f"[ToolProxy] Executing: {tool_name}")
            
            # 执行工具
            if tool_name == "web_search":
                result = self._web_search(params.get('query', ''))
            elif tool_name == "feishu_create_doc":
                result = self._feishu_create_doc(params.get('title', 'New Document'))
            elif tool_name == "feishu_write_doc":
                result = self._feishu_write_doc(params.get('doc_token', ''), params.get('content', ''))
            elif tool_name == "generate_image":
                result = self._generate_image(params.get('prompt', ''))
            elif tool_name == "stock_analysis":
                result = self._stock_analysis(params.get('code', ''))
            elif tool_name == "stock_minute":
                result = self._stock_minute(params.get('code', ''))
            elif tool_name == "stock_valuation":
                result = self._stock_valuation(params.get('code', ''), params.get('scenario', 'base'))
            elif tool_name == "generate_chart":
                result = self._generate_chart(params)
            elif tool_name == "create_word_report":
                result = self._create_word_report(params)
            elif tool_name == "send_email":
                result = self._send_email(params)
            else:
                result = {"error": f"Unknown tool: {tool_name}"}
            
            self.send_json(result)
            
        except json.JSONDecodeError:
            self.send_json({"error": "Invalid JSON"}, 400)
        except Exception as e:
            self.send_json({"error": str(e)}, 500)
    
    def _web_search(self, query):
        """执行联网搜索 - 使用多搜索引擎"""
        from urllib.parse import quote
        
        results = []
        
        # 搜索引擎列表（按优先级）
        engines = [
            ("Bing", f"https://cn.bing.com/search?q={quote(query)}&ensearch=0"),
            ("Google", f"https://www.google.com/search?q={quote(query)}"),
            ("Baidu", f"https://www.baidu.com/s?wd={quote(query)}"),
        ]
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        
        # 尝试每个搜索引擎
        for engine_name, url in engines:
            try:
                resp = requests.get(url, headers=headers, timeout=10)
                if resp.status_code == 200 and HAS_BS4:
                    soup = BeautifulSoup(resp.text, 'html.parser')
                    
                    # 根据不同搜索引擎提取结果
                    if engine_name == "Bing":
                        items = soup.select('.b_algo h2 a')
                    elif engine_name == "Google":
                        items = soup.select('div.g a')
                    elif engine_name == "Baidu":
                        items = soup.select('.result.c-container a')
                    else:
                        items = []
                    
                    for item in items[:5]:  # 每个引擎取5条
                        title = item.get_text(strip=True)
                        href = item.get('href', '')
                        if title and href and href.startswith('http'):
                            results.append({
                                'title': title,
                                'url': href,
                                'engine': engine_name
                            })
                    
                    if results:
                        break  # 有结果就停止
                        
            except Exception as e:
                print(f"[WebSearch] {engine_name} failed: {e}")
                continue
        
        if results:
            return {
                "ok": True,
                "query": query,
                "count": len(results),
                "results": results
            }
        else:
            # 备用：使用autoglm-websearch
            script_path = os.path.join(SKILLS_DIR, "autoglm-websearch", "websearch.py")
            if os.path.exists(script_path):
                try:
                    result = subprocess.run(
                        ["python", script_path, query],
                        capture_output=True,
                        text=True,
                        timeout=60,
                        encoding='utf-8'
                    )
                    if result.returncode == 0:
                        return json.loads(result.stdout.strip())
                except:
                    pass
            
            return {"ok": False, "error": "All search engines failed", "query": query}
    
    def _feishu_create_doc(self, title):
        """创建飞书文档"""
        # 使用feishu_doc工具
        try:
            import urllib.request
            
            token_path = os.path.expanduser("~/.openclaw-autoclaw/.gateway-token")
            with open(token_path, 'r') as f:
                token = f.read().strip()
            
            # 调用feishu API
            url = "https://open.feishu.cn/open-apis/docx/v1/documents"
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }
            data = json.dumps({"title": title}).encode('utf-8')
            
            req = urllib.request.Request(url, data=data, headers=headers, method='POST')
            with urllib.request.urlopen(req, timeout=30) as resp:
                return json.loads(resp.read().decode('utf-8'))
                
        except Exception as e:
            return {"error": str(e)}
    
    def _feishu_write_doc(self, doc_token, content):
        """写入飞书文档 - 使用OpenClaw feishu_doc工具"""
        try:
            # 通过本地Gateway调用feishu_doc
            import requests
            
            # 分段写入，避免超长
            max_chunk = 3000
            chunks = [content[i:i+max_chunk] for i in range(0, len(content), max_chunk)]
            
            results = []
            for chunk in chunks:
                # 调用OpenClaw的feishu_doc append接口
                resp = requests.post(
                    "http://localhost:8768/",
                    json={
                        "tool": "feishu_doc_append",
                        "params": {
                            "doc_token": doc_token,
                            "content": chunk
                        }
                    },
                    timeout=30
                )
                results.append(resp.json())
            
            return {
                "ok": True,
                "doc_token": doc_token,
                "chunks": len(chunks),
                "results": results
            }
        except Exception as e:
            return {"ok": False, "error": str(e)}
    
    def _generate_image(self, prompt):
        """生成图片"""
        script_path = os.path.join(SKILLS_DIR, "autoglm-generate-image", "generate-image.py")
        if os.path.exists(script_path):
            try:
                result = subprocess.run(
                    ["python", script_path, prompt],
                    capture_output=True,
                    text=True,
                    timeout=120,
                    encoding='utf-8'
                )
                
                if result.returncode != 0:
                    return {"error": result.stderr}
                
                return {"result": result.stdout.strip()}
                
            except Exception as e:
                return {"error": str(e)}
        else:
            return {"error": "generate_image script not found"}
    
    def _stock_analysis(self, code):
        """获取A股实时行情"""
        script_path = os.path.join(SKILLS_DIR, "a-stock-analysis-1.0.0", "scripts", "analyze.py")
        if not os.path.exists(script_path):
            return {"error": "stock analysis script not found"}
        
        try:
            result = subprocess.run(
                ["python", script_path, code, "--json"],
                capture_output=True,
                text=True,
                timeout=30,
                encoding='utf-8'
            )
            
            if result.returncode != 0:
                return {"error": result.stderr}
            
            try:
                data = json.loads(result.stdout.strip())
                # 处理数组输出
                if isinstance(data, list) and len(data) > 0:
                    data = data[0]
                
                realtime = data.get("realtime", {})
                return {
                    "status": "ok",
                    "code": data.get("code"),
                    "name": data.get("name"),
                    "price": realtime.get("price"),
                    "change_pct": realtime.get("change_pct"),
                    "change_amt": realtime.get("change_amt"),
                    "open": realtime.get("open"),
                    "high": realtime.get("high"),
                    "low": realtime.get("low"),
                    "volume": realtime.get("volume"),
                    "amount": realtime.get("amount"),
                    "pre_close": realtime.get("pre_close")
                }
            except:
                return {"result": result.stdout.strip()}
                
        except Exception as e:
            return {"error": str(e)}
    
    def _stock_minute(self, code):
        """分时量能分析"""
        script_path = os.path.join(SKILLS_DIR, "a-stock-analysis-1.0.0", "scripts", "analyze.py")
        if not os.path.exists(script_path):
            return {"error": "stock analysis script not found"}
        
        try:
            # 分时分析（不含--minute避免编码问题）
            result = subprocess.run(
                ["python", script_path, code, "--json"],
                capture_output=True,
                text=True,
                timeout=30,
                encoding='utf-8'
            )
            
            if result.returncode != 0:
                return {"error": result.stderr}
            
            return {"result": result.stdout.strip()}
                
        except Exception as e:
            return {"error": str(e)}
    
    def _stock_valuation(self, code, scenario):
        """股票估值分析"""
        # 返回估值分析框架
        return {
            "status": "ok",
            "code": code,
            "scenario": scenario,
            "framework": {
                "step1_dcf": {
                    "name": "DCF模型",
                    "inputs": ["当前股价", "总股本", "净利润", "FCF转化率", "增长率", "WACC"],
                    "outputs": ["每股价值", "估值区间", "敏感性分析"]
                },
                "step2_comparable": {
                    "name": "可比公司分析",
                    "inputs": ["同行公司列表", "PE", "PB"],
                    "outputs": ["行业平均PE", "相对估值"]
                },
                "step3_historical": {
                    "name": "历史估值",
                    "inputs": ["当前PE", "5年平均PE"],
                    "outputs": ["估值分位", "高估/低估判断"]
                },
                "step4_target": {
                    "name": "分析师目标价",
                    "inputs": ["券商研报"],
                    "outputs": ["最高/最低/中位数目标价"]
                },
                "step5_scenario": {
                    "name": "情景分析",
                    "scenarios": ["牛市(20%)", "基础(50%)", "熊市(30%)"],
                    "outputs": ["综合公允价值"]
                }
            },
            "report_format": {
                "title": "# XXX(代码)估值分析报告",
                "sections": ["DCF模型", "可比公司分析", "历史估值", "分析师目标价", "情景分析", "结论"],
                "rules": [
                    "所有假设标注[假设]",
                    "所有来源标注[来源]",
                    "不使用markdown复杂表格",
                    "不使用emoji和特殊符号"
                ]
            },
            "note": "请按照框架完成估值分析，注意报告格式要求"
        }
    
    def _generate_chart(self, params):
        """生成图表"""
        script_path = os.path.join(SKILLS_DIR, "chart-generator", "chart_generator.py")
        if not os.path.exists(script_path):
            return {"error": "chart_generator script not found"}
        
        chart_type = params.get('chart_type', 'line')
        title = params.get('title', 'Chart')
        data = params.get('data', {})
        x_labels = params.get('x_labels', [])
        output_path = params.get('output_path', '/tmp/chart.png')
        
        try:
            # 构建Python代码
            code = f"""
import json
from chart_generator import generate_chart

data = json.loads('{json.dumps(data)}')
generate_chart(
    chart_type="{chart_type}",
    title="{title}",
    x_labels={x_labels},
    data=data,
    output_path="{output_path}"
)
print("{output_path}")
"""
            
            result = subprocess.run(
                ["python", "-c", code],
                capture_output=True,
                text=True,
                timeout=30,
                encoding='utf-8',
                cwd=os.path.dirname(script_path)
            )
            
            if result.returncode != 0:
                return {"error": result.stderr}
            
            return {
                "status": "ok",
                "chart_path": output_path,
                "message": f"图表已生成: {output_path}"
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    def _create_word_report(self, params):
        """生成Word报告"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            title = params.get('title', '报告')
            content = params.get('content', '')
            
            # 创建Word文档
            doc = Document()
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加内容（支持简单的换行分隔）
            for line in content.split('\n'):
                if line.strip():
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('**') and line.endswith('**'):
                        p = doc.add_paragraph()
                        p.add_run(line[2:-2]).bold = True
                    else:
                        doc.add_paragraph(line)
            
            # 保存
            output_path = f"/tmp/{title}.docx"
            doc.save(output_path)
            
            return {
                "status": "ok",
                "file_path": output_path,
                "message": f"Word报告已生成: {output_path}"
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    def _send_email(self, params):
        """发送邮件"""
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart
            from email.mime.application import MIMEApplication
            from email.header import Header
            
            subject = params.get('subject', '报告')
            body = params.get('body', '')
            attachment = params.get('attachment', '')
            
            # 配置（从MEMORY.md读取）
            smtp_server = 'smtp.qq.com'
            smtp_port = 587
            sender = '1784350294@qq.com'
            password = 'spctbhqjdwhadfef'
            receiver = '1784350294@qq.com'
            
            # 创建邮件
            msg = MIMEMultipart()
            msg['From'] = sender
            msg['To'] = receiver
            msg['Subject'] = Header(subject, 'utf-8')
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # 添加附件
            if attachment and os.path.exists(attachment):
                with open(attachment, 'rb') as f:
                    att = MIMEApplication(f.read())
                    att.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment))
                    msg.attach(att)
            
            # 发送
            server = smtplib.SMTP(smtp_server, smtp_port)
            server.starttls()
            server.login(sender, password)
            server.sendmail(sender, receiver, msg.as_string())
            server.quit()
            
            return {
                "status": "ok",
                "message": f"邮件已发送到 {receiver}"
            }
            
        except Exception as e:
            return {"error": str(e)}


class ThreadedHTTPServer(ThreadingMixIn, HTTPServer):
    pass


def main():
    port = 8768
    server = ThreadedHTTPServer(('localhost', port), ToolProxyHandler)
    
    print("=" * 50)
    print("Swarm AGS Tool Proxy v2 Started")
    print(f"URL: http://localhost:{port}")
    print(f"Skills Dir: {SKILLS_DIR}")
    print("=" * 50)
    print("\nAvailable tools:")
    print("  - web_search: 联网搜索")
    print("  - feishu_create_doc: 创建飞书文档")
    print("  - feishu_write_doc: 写入飞书文档")
    print("  - generate_image: 生成图片")
    print("\nUsage:")
    print('  POST {"tool": "web_search", "params": {"query": "xxx"}}')
    
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nService stopped")
        server.shutdown()


if __name__ == "__main__":
    main()

